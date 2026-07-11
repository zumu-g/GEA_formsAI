"""Rendering pipeline (U4).

Programmatically fills the .docx template using python-docx (the template has no
named text fields — values live in indexed label/value table cells), ticks the
correct legacy form checkboxes, writes the DOCX, then renders a PDF via
LibreOffice headless (``soffice``). PDF is the primary deliverable; the DOCX is
always kept. If LibreOffice is unavailable the DOCX is still produced and a
warning is returned rather than failing the whole fill.

``compute_blank_fields`` walks the spec's declared fields against the resolved
context (KTD6) — this is the single source of truth for the blank/filled report.
"""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

from .errors import RenderError
from .formspec import FormSpec


def compute_blank_fields(spec: FormSpec, context: dict[str, str]) -> list[str]:
    """Declared text fields that resolved empty (selectors excluded)."""

    selectors = set(spec.selector_fields)
    return [
        f
        for f in spec.declared_fields
        if f not in selectors and not str(context.get(f, "")).strip()
    ]


def render(
    spec: FormSpec,
    context: dict[str, str],
    out_dir: str | Path,
    *,
    basename: str | None = None,
) -> tuple[Path | None, Path | None, list[str]]:
    """Render the form. Returns ``(docx_path_or_None, pdf_path_or_None, warnings)``.

    Dispatches on ``spec.engine``: "docx" fills the Word template below;
    "pdf_overlay" stamps a scanned-PDF template (no DOCX output).
    """

    if spec.engine == "pdf_overlay":
        from .render_overlay import render_overlay

        return render_overlay(spec, context, out_dir, basename=basename)

    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    name = basename or spec.key
    docx_path = out / f"{name}.docx"
    warnings: list[str] = []

    document = docx.Document(str(spec.template))
    _apply_text_ops(spec, document, context)
    _apply_checkbox_ops(spec, document, context)
    document.save(str(docx_path))

    pdf_path = _to_pdf(docx_path, out)
    if pdf_path is None:
        hint = (
            "brew install --cask libreoffice"
            if sys.platform == "darwin"
            else "install LibreOffice and ensure `soffice` is on PATH"
        )
        warnings.append(
            "PDF skipped — LibreOffice (soffice) not found on PATH; DOCX produced only. "
            f"Fix: {hint}"
        )

    return docx_path, pdf_path, warnings


def _apply_text_ops(spec: FormSpec, document, context: dict[str, str]) -> None:
    tables = document.tables
    for op in spec.text_ops:
        value = str(context.get(op.field_name, "") or "")
        if not value:
            continue
        if op.table_index >= len(tables):
            raise RenderError(
                f"template has no table {op.table_index} for field {op.field_name!r}"
            )
        table = tables[op.table_index]
        try:
            cell = table.rows[op.row_index].cells[op.cell_index]
        except IndexError as exc:
            raise RenderError(
                f"template cell out of range for field {op.field_name!r} "
                f"(table {op.table_index}, row {op.row_index}, cell {op.cell_index})"
            ) from exc
        _set_cell_text(cell, value)


def _set_cell_text(cell, value: str) -> None:
    """Write value into the first paragraph's first run, preserving formatting."""

    # Some templates ship value cells with no paragraph element at all
    # (e.g. the Allmain letter of offer) — add one rather than crashing.
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if para.runs:
        para.runs[0].text = value
        for extra in para.runs[1:]:
            extra.text = ""
    else:
        para.add_run(value)
    _allow_cell_to_grow(cell)


def _allow_cell_to_grow(cell) -> None:
    """Let a filled cell wrap and its row grow, instead of clipping long values.

    CAV's own templates format every fillable cell as single-line (``w:noWrap``)
    inside a fixed-height row (``w:trHeight hRule="exact"``) — sized for a short
    answer. An engine-computed value that runs longer (an address, a full legal
    reason string) gets visually cropped rather than wrapping. Since the form is
    still being filled, the row has no visible content yet to reflow around, so
    switching the row to ``atLeast`` (a minimum, not a cap) only ever grows it.
    """

    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        no_wrap = tcPr.find(qn("w:noWrap"))
        if no_wrap is not None:
            tcPr.remove(no_wrap)

    tr = tc.getparent()
    trPr = tr.find(qn("w:trPr"))
    tr_height = trPr.find(qn("w:trHeight")) if trPr is not None else None
    if tr_height is not None and tr_height.get(qn("w:hRule")) == "exact":
        tr_height.set(qn("w:hRule"), "atLeast")


def _apply_checkbox_ops(spec: FormSpec, document, context: dict[str, str]) -> None:
    tables = document.tables
    for op in spec.checkbox_ops:
        selector = str(context.get(op.selector_field, "") or "")
        position = op.options.get(selector)
        if position is None:
            continue  # no/unknown period → leave all boxes unticked
        if op.table_index >= len(tables):
            raise RenderError(
                f"template has no table {op.table_index} for checkbox "
                f"{op.selector_field!r}"
            )
        checkboxes = tables[op.table_index]._tbl.findall(".//" + qn("w:checkBox"))
        if position >= len(checkboxes):
            raise RenderError(
                f"checkbox position {position} out of range in table {op.table_index}"
            )
        _set_checkbox(checkboxes[position], True)


def _set_checkbox(checkbox_el, checked: bool) -> None:
    """Set a legacy form checkbox's checked state in the OOXML.

    Replaces ``<w:default>`` and ensures a ``<w:checked w:val="1">`` so both Word
    and LibreOffice render the tick on open.
    """

    val = "1" if checked else "0"
    # default
    default = checkbox_el.find(qn("w:default"))
    if default is None:
        default = checkbox_el.makeelement(qn("w:default"), {})
        checkbox_el.append(default)
    default.set(qn("w:val"), val)
    # checked
    checked_el = checkbox_el.find(qn("w:checked"))
    if checked_el is None:
        checked_el = checkbox_el.makeelement(qn("w:checked"), {})
        checkbox_el.insert(0, checked_el)
    checked_el.set(qn("w:val"), val)


def _to_pdf(docx_path: Path, out_dir: Path) -> Path | None:
    """Convert DOCX → PDF via LibreOffice headless. Returns None if unavailable."""

    soffice = _find_soffice()
    if soffice is None:
        return None
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=120,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
        raise RenderError(f"LibreOffice PDF conversion failed: {exc}") from exc
    pdf_path = out_dir / f"{docx_path.stem}.pdf"
    return pdf_path if pdf_path.exists() else None


def _find_soffice() -> str | None:
    for candidate in ("soffice", "libreoffice"):
        found = shutil.which(candidate)
        if found:
            return found
    mac = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    return str(mac) if mac.exists() else None
