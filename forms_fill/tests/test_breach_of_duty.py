import docx
import pytest

from forms_fill.forms.breach_of_duty_notice.spec import SPEC, build_context
from forms_fill.models import TenancyBundle
from forms_fill.render import render

DUTY_PROVISIONS = ("60(1)", "60(2)", "61", "63(1)", "63A", "70(2)", "70(3)", "89")


def _bundle(d):
    return TenancyBundle.model_validate(d)


@pytest.fixture
def breach_caller_fields():
    return {
        "breach_details": "60(1)-nuisance: you have used the premises in a way that causes a nuisance.",
        "loss_or_damage": "None claimed.",
        "remedy_action": "Cease the nuisance behaviour immediately.",
        "compensation_amount": "0",
    }


def test_context_has_premises_renter_provider(sample_bundle_dict, breach_caller_fields):
    ctx = build_context(_bundle(sample_bundle_dict), breach_caller_fields)
    assert ctx["premises_address"] == "12 Example Street, Richmond"
    assert ctx["renter1_name"] == "Jane Alice Smith"
    assert ctx["provider_name"] == "Robert James Owner"


@pytest.mark.parametrize("provision", DUTY_PROVISIONS)
def test_caller_fields_render_verbatim_for_each_provision(sample_bundle_dict, provision):
    fields = {
        "breach_details": f"{provision}: breach description text.",
        "loss_or_damage": "N/A",
        "remedy_action": "Remedy the breach.",
        "compensation_amount": "150",
    }
    ctx = build_context(_bundle(sample_bundle_dict), fields)
    assert ctx["breach_details"] == f"{provision}: breach description text."
    assert ctx["loss_or_damage"] == "N/A"
    assert ctx["remedy_action"] == "Remedy the breach."
    assert ctx["compensation_amount"] == "150"


def test_missing_caller_field_renders_blank_not_guessed(sample_bundle_dict):
    ctx = build_context(_bundle(sample_bundle_dict), {})
    assert ctx["breach_details"] == ""
    assert ctx["loss_or_damage"] == ""
    assert ctx["remedy_action"] == ""
    assert ctx["compensation_amount"] == ""


def test_leave_delivery_and_signature_fields_not_declared():
    joined = " ".join(SPEC.declared_fields)
    for forbidden in ("delivery", "signature", "tracking"):
        assert forbidden not in joined


def test_overflow_note_for_more_than_four_renters(sample_bundle_dict, breach_caller_fields):
    renters = [{"full_name": f"Renter {i}"} for i in range(1, 6)]
    sample_bundle_dict["renters"] = renters
    ctx = build_context(_bundle(sample_bundle_dict), breach_caller_fields)
    assert ctx["renter1_name"] == "Renter 1"
    assert "extra page" in ctx["renter4_name"]
    assert "Renter 5" in ctx["renter4_name"]


def test_template_exists_and_registered():
    assert SPEC.template.exists()
    assert SPEC.key == "breach_of_duty_notice"


def test_render_writes_real_values_into_real_template_cells(
    tmp_path, sample_bundle_dict, breach_caller_fields
):
    ctx = build_context(TenancyBundle.model_validate(sample_bundle_dict), breach_caller_fields)
    docx_path, _pdf, _warn = render(SPEC, ctx, tmp_path)
    assert docx_path.exists()
    document = docx.Document(str(docx_path))
    text = " ".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "12 Example Street, Richmond" in text
    assert "Jane Alice Smith" in text
    assert "Robert James Owner" in text
    assert breach_caller_fields["breach_details"] in text
    assert breach_caller_fields["remedy_action"] in text
    # remedy_action -> table 14 row 0 cell 0; compensation_amount -> row 1 cell 1
    assert document.tables[14].rows[0].cells[0].text == breach_caller_fields["remedy_action"]
    assert (
        document.tables[14].rows[1].cells[1].text == breach_caller_fields["compensation_amount"]
    )
