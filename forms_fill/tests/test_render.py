from pathlib import Path

import docx
from docx.oxml.ns import qn

import forms_fill.render as render_mod
from forms_fill.forms.cav_rent_increase_notice.spec import SPEC, build_context
from forms_fill.models import TenancyBundle
from forms_fill.render import compute_blank_fields, render


def _ctx(sample_bundle_dict, caller_fields):
    return build_context(TenancyBundle.model_validate(sample_bundle_dict), caller_fields)


def _all_text(document) -> str:
    parts = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return " ".join(parts)


def test_render_writes_docx_with_values(tmp_path, sample_bundle_dict, caller_fields):
    ctx = _ctx(sample_bundle_dict, caller_fields)
    docx_path, _pdf, _warn = render(SPEC, ctx, tmp_path)
    assert docx_path.exists()
    text = _all_text(docx.Document(str(docx_path)))
    assert "12 Example Street, Richmond" in text
    assert "Jane Alice Smith" in text
    assert "Robert James Owner" in text
    assert "market comparison (rental CMA)" in text


def test_blank_fields_reported(tmp_path, sample_bundle_dict, caller_fields):
    # renter has no after-hours? it does in fixture; blank the provider after-hours
    sample_bundle_dict["rental_provider"]["phone_after_hours"] = ""
    ctx = _ctx(sample_bundle_dict, caller_fields)
    blanks = compute_blank_fields(SPEC, ctx)
    assert "provider_after_hours" in blanks
    # selector is never counted
    assert "rent_period" not in blanks


def test_correct_period_checkbox_ticked(tmp_path, sample_bundle_dict, caller_fields):
    caller_fields["rent_period"] = "fortnight"
    ctx = _ctx(sample_bundle_dict, caller_fields)
    docx_path, _pdf, _warn = render(SPEC, ctx, tmp_path)
    document = docx.Document(str(docx_path))
    # table 15 (current rent): boxes are week(0), fortnight(1), calendar month(2)
    boxes = document.tables[15]._tbl.findall(".//" + qn("w:checkBox"))
    checked = [
        b.find(qn("w:checked")) is not None
        and b.find(qn("w:checked")).get(qn("w:val")) == "1"
        for b in boxes
    ]
    assert checked == [False, True, False]


def test_filled_and_blank_counts_consistent(tmp_path, sample_bundle_dict, caller_fields):
    ctx = _ctx(sample_bundle_dict, caller_fields)
    blanks = compute_blank_fields(SPEC, ctx)
    selectors = set(SPEC.selector_fields)
    declared_text = [f for f in SPEC.declared_fields if f not in selectors]
    # every declared text field is either filled or blank, no overlap
    filled = [f for f in declared_text if f not in blanks]
    assert len(filled) + len(blanks) == len(declared_text)


def test_pdf_skipped_warning_names_fix(tmp_path, sample_bundle_dict, caller_fields, monkeypatch):
    monkeypatch.setattr(render_mod, "_find_soffice", lambda: None)
    monkeypatch.setattr(render_mod.sys, "platform", "darwin")
    ctx = _ctx(sample_bundle_dict, caller_fields)
    _docx, pdf_path, warn = render(SPEC, ctx, tmp_path)
    assert pdf_path is None
    assert warn == [
        "PDF skipped — LibreOffice (soffice) not found on PATH; DOCX produced only. "
        "Fix: brew install --cask libreoffice"
    ]

    monkeypatch.setattr(render_mod.sys, "platform", "linux")
    _docx, pdf_path, warn = render(SPEC, ctx, tmp_path, basename="rent_increase_linux")
    assert warn == [
        "PDF skipped — LibreOffice (soffice) not found on PATH; DOCX produced only. "
        "Fix: install LibreOffice and ensure `soffice` is on PATH"
    ]


def test_pdf_conversion_uses_isolated_profile_and_cleans_up(tmp_path, monkeypatch):
    # Guards the concurrency fix: each conversion must get its own throwaway
    # LibreOffice profile via -env:UserInstallation (otherwise concurrent/orphaned
    # soffice deadlock on the shared default-profile lock), and the temp profile
    # must be removed afterwards.
    captured = {}

    def fake_run(argv, **kwargs):
        captured["argv"] = argv
        captured["timeout"] = kwargs.get("timeout")
        (tmp_path / "doc.pdf").write_bytes(b"%PDF-1.4")
        return None

    monkeypatch.setattr(render_mod, "_find_soffice", lambda: "/usr/bin/soffice")
    monkeypatch.setattr(render_mod.subprocess, "run", fake_run)

    (tmp_path / "doc.docx").write_bytes(b"x")
    out = render_mod._to_pdf(tmp_path / "doc.docx", tmp_path)

    assert out == tmp_path / "doc.pdf"
    prof = [a for a in captured["argv"] if a.startswith("-env:UserInstallation=file://")]
    assert len(prof) == 1, "each conversion needs its own isolated profile"
    assert captured["timeout"] and captured["timeout"] > 0, "must be bounded"
    profile_path = prof[0].split("file://", 1)[1]
    assert not Path(profile_path).exists(), "temp profile must be cleaned up"


def test_filled_cells_allow_wrap_and_growth(tmp_path, sample_bundle_dict, caller_fields):
    # CAV's own templates format every fillable cell single-line (w:noWrap) inside
    # a fixed-height row (hRule="exact") -- sized for a short answer. A longer
    # engine-computed value (e.g. a two-line address) must not render cropped.
    ctx = _ctx(sample_bundle_dict, caller_fields)
    docx_path, _pdf, _warn = render(SPEC, ctx, tmp_path)
    document = docx.Document(str(docx_path))
    # table 1 cell 0 = premises_address; table 10 cell 0 = provider_name -- both
    # get real text written into them by build_context/TEXT_OPS.
    for table_index in (1, 10):
        cell = document.tables[table_index].rows[0].cells[0]
        tcPr = cell._tc.find(qn("w:tcPr"))
        assert tcPr.find(qn("w:noWrap")) is None, f"table {table_index} cell still noWrap"
        trPr = document.tables[table_index].rows[0]._tr.find(qn("w:trPr"))
        tr_height = trPr.find(qn("w:trHeight"))
        assert tr_height.get(qn("w:hRule")) == "atLeast", (
            f"table {table_index} row still fixed-height (would clip a long value)"
        )


def test_merged_cell_ops_append_rather_than_overwrite(tmp_path):
    """Two text ops resolving to one vertically-merged cell keep both values
    (regression: provider_name was lost under provider_company_name in the
    residential rental agreements)."""
    import docx

    from forms_fill.forms.residential_rental_agreement.spec import SPEC
    from forms_fill.render import _apply_text_ops

    document = docx.Document(str(SPEC.template))
    ctx = {"provider_name": "Owner Person", "provider_company_name": "Owner Co Pty Ltd"}
    _apply_text_ops(SPEC, document, ctx)
    text = document.tables[3].rows[0].cells[1].text
    assert "Owner Person" in text and "Owner Co Pty Ltd" in text


def test_pdf_false_skips_conversion_and_warning(tmp_path):
    from forms_fill.core import fill_form
    from forms_fill.models import FillRequest

    result = fill_form(
        FillRequest(
            form="general_notice",
            identifiers={"lot_id": "mp_xyz789", "tenancy_id": "lease_abc123"},
            fields={"reason_for_notice": "s55 test"},
            out_dir=str(tmp_path),
            provider="fixture",
            pdf=False,
        )
    )
    assert result.ok
    assert result.files.docx and result.files.pdf is None
    assert not any("LibreOffice" in w for w in result.warnings)
