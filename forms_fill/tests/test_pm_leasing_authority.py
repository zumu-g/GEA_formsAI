"""Tests for pm_exclusive_leasing_authority (plan 2026-07-29-001, U3)."""

from __future__ import annotations

import docx

from forms_fill.forms.pm_exclusive_leasing_authority.spec import SPEC, build_context
from forms_fill.registry import available_forms, form_catalogue, get_form_spec
from forms_fill.render import compute_blank_fields, render

FIELDS = {
    "owner_names": "Jane Owner",
    "owner_address": "1 Owner St, Berwick VIC 3806",
    "owner_phone": "0400 000 000",
    "owner_email": "jane@example.com",
    "property_address": "7 Example Avenue, Berwick VIC 3806",
    "exclusive_leasing_days": "90",
    "continuing_leasing_days": "90",
    "fixed_management_period": "n/a",
    "rent_per_week": "450",
    "security_bond": "1950",
    "urgent_repair_limit": "2500",
    "leasing_fee": "1.5 weeks rent",
    "releasing_fee": "$165",
    "managing_fee": "6.6%",
    "commission_sharing": "No",
    "rebate_entitlement": "No",
    "agreement_date": "2026-07-29",
}


def _all_text(document) -> str:
    parts = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return " ".join(parts)


def test_registered_in_registry():
    assert "pm_exclusive_leasing_authority" in available_forms()
    assert get_form_spec("pm_exclusive_leasing_authority") is SPEC


def test_catalogue_entry_is_gea_pm():
    entry = next(
        f for f in form_catalogue() if f["key"] == "pm_exclusive_leasing_authority"
    )
    assert entry["category"] == "GEA PM"
    assert entry["short_title"] == "Exclusive Leasing Authority"


def test_caller_fields_render_verbatim(tmp_path):
    ctx = build_context(None, FIELDS)
    docx_path, _pdf, _warn = render(SPEC, ctx, tmp_path, pdf=False)
    text = _all_text(docx.Document(str(docx_path)))
    assert "Jane Owner" in text
    assert "7 Example Avenue, Berwick VIC 3806" in text
    assert "1.5 weeks rent" in text


def test_agency_defaults_fill_agent_block(tmp_path):
    ctx = build_context(None, FIELDS)
    docx_path, _pdf, _warn = render(SPEC, ctx, tmp_path, pdf=False)
    text = _all_text(docx.Document(str(docx_path)))
    assert "Grants Estate Agents" in text
    assert "Stuart Grant" in text


def test_blank_fields_reported_for_missing_input():
    ctx = build_context(None, {})
    blanks = compute_blank_fields(SPEC, ctx)
    assert "property_address" in blanks
    # agency defaults are always populated, never blank
    assert "agent_name" not in blanks
