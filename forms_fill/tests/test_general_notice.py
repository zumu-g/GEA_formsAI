import docx
import pytest

from forms_fill.forms.general_notice.spec import SPEC, build_context
from forms_fill.models import TenancyBundle
from forms_fill.render import render

REASONS = ("55(1)", "66(2)", "78(1)", "79(1)", "79(2)", "91M")


def _bundle(d):
    return TenancyBundle.model_validate(d)


@pytest.fixture
def general_caller_fields():
    return {"reason_for_notice": "55(1): I have paid utility charges of $120 that are your responsibility."}


def test_context_has_premises_renter_provider(sample_bundle_dict, general_caller_fields):
    ctx = build_context(_bundle(sample_bundle_dict), general_caller_fields)
    assert ctx["premises_address"] == "12 Example Street, Richmond"
    assert ctx["renter1_name"] == "Jane Alice Smith"
    assert ctx["provider_name"] == "Robert James Owner"


@pytest.mark.parametrize("reason", REASONS)
def test_caller_field_renders_verbatim_for_each_reason(sample_bundle_dict, reason):
    fields = {"reason_for_notice": f"{reason}: notice text for this ground."}
    ctx = build_context(_bundle(sample_bundle_dict), fields)
    assert ctx["reason_for_notice"] == f"{reason}: notice text for this ground."


def test_missing_caller_field_renders_blank_not_guessed(sample_bundle_dict):
    ctx = build_context(_bundle(sample_bundle_dict), {})
    assert ctx["reason_for_notice"] == ""


def test_leave_delivery_and_signature_fields_not_declared():
    joined = " ".join(SPEC.declared_fields)
    for forbidden in ("delivery", "signature", "tracking"):
        assert forbidden not in joined


def test_overflow_note_for_more_than_four_renters(sample_bundle_dict, general_caller_fields):
    renters = [{"full_name": f"Renter {i}"} for i in range(1, 6)]
    sample_bundle_dict["renters"] = renters
    ctx = build_context(_bundle(sample_bundle_dict), general_caller_fields)
    assert ctx["renter1_name"] == "Renter 1"
    assert "extra page" in ctx["renter4_name"]
    assert "Renter 5" in ctx["renter4_name"]


def test_template_exists_and_registered():
    assert SPEC.template.exists()
    assert SPEC.key == "general_notice"


def test_render_writes_real_values_into_real_template_cells(
    tmp_path, sample_bundle_dict, general_caller_fields
):
    ctx = build_context(TenancyBundle.model_validate(sample_bundle_dict), general_caller_fields)
    docx_path, _pdf, _warn = render(SPEC, ctx, tmp_path)
    assert docx_path.exists()
    document = docx.Document(str(docx_path))
    text = " ".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "12 Example Street, Richmond" in text
    assert "Jane Alice Smith" in text
    assert "Robert James Owner" in text
    assert general_caller_fields["reason_for_notice"] in text
    cell_text = document.tables[11].rows[0].cells[0].text.strip()
    assert cell_text == general_caller_fields["reason_for_notice"]
