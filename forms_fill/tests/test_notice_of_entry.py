import docx
import pytest

from forms_fill.forms.notice_of_entry.spec import SPEC, build_context
from forms_fill.models import TenancyBundle
from forms_fill.render import render

REASONS = (
    "86(1)(a) - inspection prospective renter",
    "86(1)(c) - legal duty",
    "86(1)(d) - valuation",
    "86(1)(e) - breach of duty",
    "86(1)(g) - family or personal violence",
)


def _bundle(d):
    return TenancyBundle.model_validate(d)


@pytest.fixture
def entry_caller_fields():
    return {
        "entry_date": "2026-08-01",
        "entry_window": "9am to 5pm",
        "entry_reason": "86(1)(a): I have given you a written offer to sell the premises.",
    }


def test_context_has_premises_renter_provider(sample_bundle_dict, entry_caller_fields):
    ctx = build_context(_bundle(sample_bundle_dict), entry_caller_fields)
    assert ctx["premises_address"] == "12 Example Street, Richmond"
    assert ctx["renter1_name"] == "Jane Alice Smith"
    assert ctx["provider_name"] == "Robert James Owner"


@pytest.mark.parametrize("reason", REASONS)
def test_caller_field_renders_verbatim_for_each_reason(sample_bundle_dict, reason):
    fields = {"entry_date": "2026-08-01", "entry_window": "9am to 5pm", "entry_reason": reason}
    ctx = build_context(_bundle(sample_bundle_dict), fields)
    assert ctx["entry_reason"] == reason


def test_missing_caller_field_renders_blank_not_guessed(sample_bundle_dict):
    ctx = build_context(_bundle(sample_bundle_dict), {})
    assert ctx["entry_date"] == ""
    assert ctx["entry_window"] == ""
    assert ctx["entry_reason"] == ""


def test_leave_delivery_and_signature_fields_not_declared():
    joined = " ".join(SPEC.declared_fields)
    for forbidden in ("delivery", "signature", "tracking"):
        assert forbidden not in joined


def test_overflow_note_for_more_than_four_renters(sample_bundle_dict, entry_caller_fields):
    renters = [{"full_name": f"Renter {i}"} for i in range(1, 6)]
    sample_bundle_dict["renters"] = renters
    ctx = build_context(_bundle(sample_bundle_dict), entry_caller_fields)
    assert ctx["renter1_name"] == "Renter 1"
    assert "extra page" in ctx["renter4_name"]
    assert "Renter 5" in ctx["renter4_name"]


def test_template_exists_and_registered():
    assert SPEC.template.exists()
    assert SPEC.key == "notice_of_entry"


def test_render_writes_real_values_into_real_template_cells(
    tmp_path, sample_bundle_dict, entry_caller_fields
):
    ctx = build_context(TenancyBundle.model_validate(sample_bundle_dict), entry_caller_fields)
    docx_path, _pdf, _warn = render(SPEC, ctx, tmp_path)
    assert docx_path.exists()
    document = docx.Document(str(docx_path))
    text = " ".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "12 Example Street, Richmond" in text
    assert "Jane Alice Smith" in text
    assert "Robert James Owner" in text
    assert entry_caller_fields["entry_reason"] in text
    assert document.tables[10].rows[0].cells[1].text == entry_caller_fields["entry_date"]
    assert document.tables[10].rows[1].cells[1].text == entry_caller_fields["entry_window"]
    assert document.tables[12].rows[0].cells[0].text == entry_caller_fields["entry_reason"]
